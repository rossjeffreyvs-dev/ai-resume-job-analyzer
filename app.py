import os
from dotenv import load_dotenv
from openai import OpenAI
from flask import Flask, request, render_template
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document

# ==========================================================
#  Environment Setup
# ==========================================================
# Load .env for local development; Lightsail provides env vars directly
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# ==========================================================
#  Initialize OpenAI Client
# ==========================================================
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not OPENAI_API_KEY:
    raise RuntimeError(
        "❌ OPENAI_API_KEY is not set. "
        "For local dev: create a .env file with your API key.\n"
        "For production: add OPENAI_API_KEY in Lightsail container environment variables."
    )

client = OpenAI(api_key=OPENAI_API_KEY)

# ==========================================================
#  Helper Functions
# ==========================================================
def extract_text_from_file(filepath):
    """Extract text from a PDF or DOCX file."""
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()

    if ext == ".pdf":
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif ext == ".docx":
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def generate_match_response(job_description, resume_text):
    """Call OpenAI API to compare resume and job description."""
    prompt = f"""
You are a career assistant. A job description and a resume are provided.
Return a clear, bullet-point comparison of how well the resume matches or does not match the job description.
Use phrases like "✔️ Matches" and "❌ Missing" to explain each point.

Job Description:
{job_description}

Resume:
{resume_text}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=1200
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"⚠️ Error generating match response: {e}"

# ==========================================================
#  Routes
# ==========================================================
@app.route('/', methods=['GET', 'POST'])
def index():
    match_report = None
    error = None

    if request.method == 'POST':
        try:
            job_desc = request.form.get('job_desc')
            resume_file = request.files.get('resume_file')

            if not job_desc or not resume_file:
                error = "Please provide both a job description and a resume file."
            else:
                filename = secure_filename(resume_file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                resume_file.save(filepath)

                resume_text = extract_text_from_file(filepath)
                match_report = generate_match_response(job_desc, resume_text)

        except Exception as e:
            error = f"⚠️ {str(e)}"

    return render_template("index.html", match_report=match_report, error=error)


# ==========================================================
#  Entry Point
# ==========================================================
if __name__ == "__main__":
    # Only used in development; Lightsail runs with Gunicorn
    print("✅ Starting Flask development server...")
    app.run(host="0.0.0.0", port=5000, debug=True)


# from dotenv import load_dotenv
# import os
# from openai import OpenAI
# from flask import Flask, request, render_template
# from werkzeug.utils import secure_filename
# from PyPDF2 import PdfReader
# from docx import Document
#
# # Load .env file for local dev
# load_dotenv()
#
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
#
# # --- OpenAI Setup ---
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
#
# if not OPENAI_API_KEY:
#     raise RuntimeError(
#         "❌ OPENAI_API_KEY is not set. "
#         "Please add it to your .env file for local dev, "
#         "or as an environment variable in your Lightsail container settings."
#     )
#
# client = OpenAI(api_key=OPENAI_API_KEY)
#

# import os
# from flask import Flask, request, render_template
# from werkzeug.utils import secure_filename
# from PyPDF2 import PdfReader
# from openai import OpenAI
# from docx import Document
# import os
#
# # --- Flask setup ---
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
#
# # --- Initialize OpenAI client from environment variable ---
# OPENAI_API_KEY="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"
# # OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
#
# if not OPENAI_API_KEY:
#     raise RuntimeError(
#         "❌ OPENAI_API_KEY is not set. "
#         "Please add it as an environment variable in your Lightsail container settings."
#     )
#
# try:
#     client = OpenAI(api_key=OPENAI_API_KEY)
#     print("✅ OpenAI client initialized successfully.")
# except Exception as e:
#     print(f"⚠️ Failed to initialize OpenAI client: {e}")
#     client = None
#
#
# # --- Helper: extract text from a PDF ---
# def extract_text_from_file(filepath):
#     """Extract text from PDF or DOCX file."""
#     _, ext = os.path.splitext(filepath)
#     ext = ext.lower()
#
#     if ext == ".pdf":
#         reader = PdfReader(filepath)
#         text = ""
#         for page in reader.pages:
#             text += page.extract_text() or ""
#         return text
#
#     elif ext == ".docx":
#         doc = Document(filepath)
#         return "\n".join([para.text for para in doc.paragraphs])
#
#     else:
#         raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")
#
#
# # --- Helper: generate AI résumé/job match report ---
# def generate_match_response(job_description, resume_text):
#     prompt = f"""
# You are a career assistant. A job description and a resume are provided.
# Return a clear, bullet-point comparison of how well the resume matches or does not match the job description.
# Use phrases like "✔️ Matches" and "❌ Missing" to explain each point.
#
# Job Description:
# {job_description}
#
# Resume:
# {resume_text}
# """
#
#     if not client:
#         return "⚠️ OpenAI client not available."
#
#     try:
#         response = client.chat.completions.create(
#             model="gpt-4o",
#             messages=[{"role": "user", "content": prompt}],
#             temperature=0.4,
#             max_tokens=1200
#         )
#         return response.choices[0].message.content.strip()
#     except Exception as e:
#         print("OpenAI API error:", e)
#         return f"⚠️ Error generating match response: {e}"
#
#
# # --- Flask route ---
# @app.route('/', methods=['GET', 'POST'])
# def index():
#     match_report = None
#
#     if request.method == 'POST':
#         job_desc = request.form.get('job_desc', '').strip()
#         resume_file = request.files.get('resume_file')
#
#         if not job_desc or not resume_file:
#             match_report = "⚠️ Please provide both a job description and a résumé."
#         else:
#             os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
#             filename = secure_filename(resume_file.filename)
#             filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             resume_file.save(filepath)
#
#             resume_text = extract_text_from_file(filepath)
#             match_report = generate_match_response(job_desc, resume_text)
#
#     return render_template("index.html", match_report=match_report)
#
#
# # --- Entry point for local dev ---
# if __name__ == "__main__":
#     port = int(os.getenv("PORT", 5000))
#     app.run(host="0.0.0.0", port=port)
